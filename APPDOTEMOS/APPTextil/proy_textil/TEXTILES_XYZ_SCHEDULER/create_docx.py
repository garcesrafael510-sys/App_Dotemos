import sys
import subprocess

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

install_and_import('docx')
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Add Title
title = document.add_heading('BITÁCORA DE PROYECTO: APP TEXTIL SCHEDULER', 0)
subtitle = document.add_paragraph('Sistema de Programación Desagregada para Producción Textil')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Introducción
document.add_heading('1. INTRODUCCIÓN Y JUSTIFICACIÓN DEL PROYECTO', level=1)

document.add_heading('1.1 El Contexto del Mercado Textil', level=2)
document.add_paragraph('La industria textil y de confección moderna opera bajo un modelo de Alta Mezcla y Bajo Volumen (High Mix / Low Volume) dictado por tendencias de moda rápidas o Fast Fashion. El mercado exige tiempos de respuesta (Lead Times) cada vez más cortos, lotes más pequeños y alta personalización. Dado que los márgenes de ganancia en confección suelen ser estrechos, la eficiencia operativa en el piso de planta no es un lujo, sino el factor determinante para la supervivencia y rentabilidad de la empresa.')

document.add_heading('1.2 El Problema Operativo (El "Por Qué")', level=2)
document.add_paragraph('Tradicionalmente, las empresas textiles planifican su piso de planta (Shop Floor) apoyándose en el conocimiento empírico del supervisor, en hojas de cálculo estáticas, o mediante intuición. Este enfoque manual es incapaz de procesar las múltiples variables dinámicas del entorno de producción de manera simultánea, tales como:')
document.add_paragraph('La variabilidad de los tiempos estándar (SAM) por cada tipo de producto y operación.', style='List Bullet')
document.add_paragraph('La restricción de la capacidad finita de la mano de obra disponible.', style='List Bullet')
document.add_paragraph('Las diferentes rutas de proceso de ensamblaje (Ensamble, Terminación, General).', style='List Bullet')

p = document.add_paragraph()
p.add_run('Consecuencias de la falta de un sistema informático:').bold = True
document.add_paragraph('1. Altos tiempos de ocio (Muda de Espera): Desbalanceo en la línea donde operarios quedan sin carga mientras otros están saturados.')
document.add_paragraph('2. Incumplimiento de fechas de entrega: Imposibilidad de proyectar con matemáticas exactas si la capacidad instalada puede absorber la carga de trabajo en el tiempo requerido.')
document.add_paragraph('3. Cuellos de botella impredecibles: Al cambiar la producción de una referencia a otra, la carga de trabajo fluctúa, generando acumulaciones de inventario en proceso (WIP) descontroladas.')

document.add_heading('1.3 La Solución Propuesta', level=2)
document.add_paragraph('Nuestra aplicación digitaliza y automatiza este proceso matemático. En lugar de un balanceo a prueba y error, el sistema provee una nivelación rápida, estandarizada y repetible mediante un algoritmo de Balanceo Heurístico Secuencial. La herramienta permite responder con exactitud: ¿Qué producir?, ¿Cuánto producir? y ¿Cuándo y con Quién producir?.')

# 2. Marco Teorico
document.add_heading('2. MARCO TEÓRICO Y CONCEPTOS CLAVE', level=1)
document.add_paragraph('Para comprender el motor de la herramienta, definimos los siguientes principios de Ingeniería de Métodos y Programación de la Producción:')
document.add_paragraph('Plan Maestro de Producción (MPS): Herramienta macro que define las cantidades agregadas a fabricar (ej. semanas o meses). Nuestra app es el paso siguiente al MPS.', style='List Bullet')
document.add_paragraph('Programación Desagregada: El enfoque central de la APP. Toma una orden específica (cantidades exactas de una Referencia X) y la desglosa al nivel más detallado: tiempo por operación (minutos) asignada a un operario específico.', style='List Bullet')
document.add_paragraph('SAM (Standard Allowed Minute / Minuto Estándar Asignado): El tiempo que un trabajador calificado necesita para realizar una operación específica a un ritmo normal. En la APP, el SAM es la "moneda" que se divide y reparte de forma equitativa entre la fuerza laboral.', style='List Bullet')
document.add_paragraph('BOM (Bill of Materials / Hoja de Ruta): A diferencia de la industria metalmecánica (lista de piezas), en confección, el BOM actúa como el "Routing" o Secuencia de Operaciones. Define el orden (Unir hombros, Pegar cuello), la máquina requerida y el SAM. Sin el BOM, el algoritmo carece de input para calcular cargas.', style='List Bullet')

document.add_heading('2.1 Celdas de Manufactura vs. Modelos Tradicionales', level=2)
document.add_paragraph('Flow Shop (Ej. Ensamblaje Automotriz): Línea continua, rígida, ideal para volúmenes masivos.', style='List Bullet')
document.add_paragraph('Job Shop (Talleres agrupados por función): Diferentes departamentos (Solo corte, solo costura plana). Genera mucho traslado y cuellos de botella impredecibles.', style='List Bullet')
document.add_paragraph('Celdas de Manufactura (Modelo de la APP): Agrupa operarios polivalentes y diversas máquinas en "mini-fábricas" dedicadas a toda la ruta de un lote. La APP asume el modelo de celda, distribuyendo la carga total entre N operarios para que trabajen como un sistema unificado.', style='List Bullet')

document.add_heading('2.2 Programación por Lotes y Tiempos de Preparación (Setup)', level=2)
document.add_paragraph('El modelo opera de manera orientada por lotes (Batch). Cambiar hilos, ajustes y adaptar la maquinaria toma tiempo (Setup Time). La programación por lotes en la APP permite agrupar todas las unidades de una misma referencia para procesarlas de corrido en la célula, diluyendo este tiempo improductivo, para luego saltar a la siguiente referencia.')

# 3. Cuellos de botella
document.add_heading('3. ANÁLISIS DE CUELLOS DE BOTELLA Y FACTORES CRÍTICOS', level=1)
document.add_paragraph('En ambientes Job Shop o Flow Shop, el cuello de botella es fácilmente visible: es la máquina física más lenta o la estación con una montaña de material esperando.')

p = document.add_paragraph()
p.add_run('¿Cómo maneja y mitiga la APP los Cuellos de Botella?\n').bold = True
p.add_run('La APP emplea el principio de "polivalencia" y "división de tareas líquidas". Si una operación requiere 90 minutos de esfuerzo y el ciclo de un operario es de 60 minutos, la operación en un modelo rígido sería un cuello de botella infranqueable. Sin embargo, nuestro algoritmo heurístico "llena" los 60 minutos del Operario 1, y asigna los 30 minutos restantes al Operario 2.')

p = document.add_paragraph()
p.add_run('Identificación del Factor Crítico en este software:\n').bold = True
p.add_run('En nuestro modelo, el cuello de botella físico individual se diluye al ser absorbido por la célula de trabajo en conjunto. El factor crítico real y que el software alerta ocurre cuando el Tiempo Total Necesario (SAM acumulado de la demanda) supera el Tiempo Total Disponible (Capacidad Finita de Operarios x Ciclo). En ese punto, el sistema entra en un déficit de "Operarios Insuficientes", señalando que la capacidad global de la planta es el límite (Restricción de Sistema).')

# 4. SIPOC
document.add_heading('4. SIPOC: SISTEMA DE PROGRAMACIÓN DE PRODUCCIÓN', level=1)
table = document.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Elemento'
hdr_cells[1].text = 'Descripción para la APP'

row_cells = table.add_row().cells
row_cells[0].text = 'S (Supplier / Proveedor)'
row_cells[1].text = '- Departamento Comercial (Ventas / Demanda)\n- Ingeniería de Producción (Métodos y Tiempos / Estudios MTM)'

row_cells = table.add_row().cells
row_cells[0].text = 'I (Input / Entrada)'
row_cells[1].text = '- Órdenes de Producción: (Referencia, Cantidad esperada).\n- Base de datos BOM / Hojas de Ruta: (Operaciones, Secuencia, SAM).\n- Restricciones de Planta: (# Operarios, Tiempo dispuesto).'

row_cells = table.add_row().cells
row_cells[0].text = 'P (Process / Procesos)'
row_cells[1].text = '1. Cálculo de Meta Takt Time (Unidades a fabricar por hora).\n2. Traducción de demanda a carga en minutos (U/H x SAM).\n3. Ejecución del Algoritmo de Balanceo Heurístico Secuencial.\n4. Diagramación de Gantt In-Memory.\n5. Consolidación de Indicadores KPIs.'

row_cells = table.add_row().cells
row_cells[0].text = 'O (Output / Salida)'
row_cells[1].text = '- Plan de Producción Desagregado y balanceado de la orden.\n- Carga laboral medida en minutos exactos por Operario.\n- Determinación de la saturación del sistema.'

row_cells = table.add_row().cells
row_cells[0].text = 'C (Customer / Cliente)'
row_cells[1].text = '- Supervisores de Planta.\n- Lideres de Célula de Manufactura.\n- Dirección y Gerencia de Producción.'


# 5. Arquitectura Técnica
document.add_heading('5. ARQUITECTURA TÉCNICA Y LÓGICA DEL MOTOR (DIAGRAMA DE FLUJO)', level=1)
document.add_paragraph('El "cerebro" de la aplicación (Motor Heurístico - scheduling_engine.py) funciona bajo los siguientes pasos metodológicos abstractos:')
document.add_paragraph('1. Lectura de Demanda: Recepción de los datos de la Orden seleccionada (Referencias y Cantidades).')
document.add_paragraph('2. Cruce Referencial: Validación estricta; el sistema ubica el BOM (SAMs) de la referencia. Si no existe, lanza una "Alerta de Falla de Ingeniería".')
document.add_paragraph('3. Cálculo de Eficacia (Meta Teórica):\nUnidades por Hora a Fabricar = (Cantidad Operarios * Tiempo de Ciclo) / Total SAM de Prenda')
document.add_paragraph('4. Desagregación Dinámica: Por cada operación en secuencia:\nTiempo Requerido por Operación = SAM de la operación * Unidades por Hora Calculadas')
document.add_paragraph('5. Bucle de Balanceo Heurístico (Asignación Tipo "Llenado de Contenedores"):')
document.add_paragraph(' El algoritmo evalúa el Operario 1 (Contenedor).', style='List Bullet 2')
document.add_paragraph(' Ingresa el tiempo de la primera operación disponible al contenedor.', style='List Bullet 2')
document.add_paragraph(' Si la operación excede el tiempo libre del operario, la capacidad libre se llena al 100% y el "excedente" de tiempo de la operación rebasa, trasladándose automáticamente a ser asignado al siguiente Operario (Contenedor 2).', style='List Bullet 2')
document.add_paragraph(' Si un operario llega a 0 minutos disponibles, el índice avanza estructuralmente al operario consecutivo.', style='List Bullet 2')
document.add_paragraph('6. Validación de Factibilidad Constante: Si la matriz de operaciones tiene remanentes (tiempos sobrantes) y el índice de operarios llega a su límite, se emite una interrupción lógica por "Capacidad Instalada Deficiente".')
document.add_paragraph('7. Motor Visual y Analítico: Representación visual gráfica tipo Gantt de asignación de minutos y despliegue del Dashboard estadístico global de balanceo.')

# 6. Conclusion
document.add_heading('6. CONCLUSIÓN', level=1)
document.add_paragraph('Esta aplicación no es únicamente un modelo de datos estructurado o CRUD, sino una herramienta computacional de Investigación de Operaciones aplicada a la industria. Conforma un ecosistema completo para tomar decisiones gerenciales informadas en un modelo desagregado, reemplazando la intuición con un algorítmico matemático puro para balanceo de capacidades en entornos celulares de alta complejidad.')

document.save('d:/Documents/APPTextil/proy_textil/TEXTILES_XYZ_SCHEDULER/BITACORA_PROYECTO.docx')
